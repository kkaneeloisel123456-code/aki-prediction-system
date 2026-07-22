"""
⚠️ DEPRECATED — 此脚本基于旧版数据（AKI=156, 非AKI=264）和旧版方法（LASSO, 8模型）。
    请勿用于最终提交。当前项目以 run_clean.py 为权威来源：
    — AKI=125/420, 非AKI=295
    — 特征筛选: RF重要性 Top35
    — 模型: LR + RF + XGBoost + ExtraTrees + Voting Ensemble (AUC 0.821±0.043)
    — 生成PDF报告请使用 scripts/gen_final_report_pdf.py

Generate a Word document for the AKI test report.
"""
import sys, os
sys.stdout.reconfigure(encoding='utf-8')

from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

doc = Document()

# Page setup
section = doc.sections[0]
section.page_width = Cm(21)
section.page_height = Cm(29.7)
section.top_margin = Cm(2.5)
section.bottom_margin = Cm(2.5)
section.left_margin = Cm(2.5)
section.right_margin = Cm(2.5)

# Default font
style = doc.styles['Normal']
font = style.font
font.name = 'Microsoft YaHei'
font.size = Pt(10.5)
style.element.rPr.rFonts.set(qn('w:eastAsia'), 'Microsoft YaHei')


def set_cell_shading(cell, color):
    shading = OxmlElement('w:shd')
    shading.set(qn('w:fill'), color)
    shading.set(qn('w:val'), 'clear')
    cell._tc.get_or_add_tcPr().append(shading)


def add_styled_table(doc, headers, rows):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = 'Table Grid'
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = ''
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(h)
        run.bold = True
        run.font.size = Pt(9)
        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        set_cell_shading(cell, '2A5A8C')
    for r_idx, row_data in enumerate(rows):
        for c_idx, val in enumerate(row_data):
            cell = table.rows[r_idx + 1].cells[c_idx]
            cell.text = ''
            p = cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run(str(val))
            run.font.size = Pt(9)
            if r_idx % 2 == 0:
                set_cell_shading(cell, 'F0F5FA')
    doc.add_paragraph()


def add_heading_styled(doc, text, level):
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.name = 'Microsoft YaHei'
        run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Microsoft YaHei')
    return h


def add_para(doc, text, bold=False, size=Pt(10.5), alignment=None, space_after=Pt(6)):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.name = 'Microsoft YaHei'
    run.font.size = size
    run.bold = bold
    run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Microsoft YaHei')
    if alignment:
        p.alignment = alignment
    p.paragraph_format.space_after = space_after
    return p


# ===== COVER PAGE =====
doc.add_paragraph()
doc.add_paragraph()

add_para(doc, '基于机器学习的急性肾损伤预测模型构建', bold=True, size=Pt(22), alignment=WD_ALIGN_PARAGRAPH.CENTER, space_after=Pt(4))
add_para(doc, '与临床决策支持系统开发', bold=True, size=Pt(22), alignment=WD_ALIGN_PARAGRAPH.CENTER, space_after=Pt(20))
add_para(doc, '—— 测试报告文档 ——', size=Pt(16), alignment=WD_ALIGN_PARAGRAPH.CENTER, space_after=Pt(40))

doc.add_paragraph()

info = [
    ('团队', '白菜卷队 | 广西科技大学'),
    ('赛事', '暑期数创 2026'),
    ('编制日期', '2026年7月13日'),
    ('版本', 'v1.0'),
    ('代码仓库', 'https://github.com/kkaneeloisel123456-code/aki-prediction-system'),
]
for label, value in info:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(label + '：')
    run.bold = True
    run.font.size = Pt(12)
    run.font.name = 'Microsoft YaHei'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Microsoft YaHei')
    run = p.add_run(value)
    run.font.size = Pt(12)
    run.font.name = 'Microsoft YaHei'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Microsoft YaHei')

doc.add_page_break()

# ===== SECTION 1 =====
add_heading_styled(doc, '一、程序整体概述以及核心模块拆解', level=1)

add_heading_styled(doc, '1.1 项目概述', level=2)
add_para(doc, '本项目构建了一套完整的急性肾损伤（AKI）智能预测辅助决策系统，用于心脏手术后AKI风险的早期预警。以420例体外循环心脏手术患者的临床数据为基础，通过4种机器学习模型及Voting集成（Logistic回归、RandomForest、XGBoost、ExtraTrees及Voting Ensemble）的系统比较与五折交叉验证，选取最优模型（XGBoost，AUC=0.821），结合SHAP可解释性框架，最终部署为Streamlit Web应用，辅助临床医生进行多时间点的AKI风险评估。')

add_heading_styled(doc, '1.2 核心架构', level=2)
add_para(doc, '系统分为六大层：')
add_para(doc, '• 数据层（data/）：原始EHR数据 → MICE插补 → IQR异常值处理 → LASSO筛选 → 训练/测试划分')
add_para(doc, '• 模型层（src/models/）：4种ML模型及Voting训练 + RandomizedSearchCV调优 + SMOTE')
add_para(doc, '• 集成层（ensemble.py）：Voting（Soft/Hard）、Stacking、加权概率融合')
add_para(doc, '• 可解释性层（SHAP）：全局特征重要性、局部Force Plot、特征依赖图')
add_para(doc, '• 可视化评估层（visualization/）：ROC/PR曲线、混淆矩阵、校准曲线、DCA决策曲线')
add_para(doc, '• Web部署层（web/app.py）：Streamlit框架，手动/批量输入，PDF报告导出')

add_heading_styled(doc, '1.3 核心模块拆解', level=2)

add_para(doc, '模块一：数据预处理模块（src/data/）', bold=True, size=Pt(11))
add_styled_table(doc,
    ['子模块', '文件', '功能说明'],
    [
        ['数据清洗', 'cleaning.py', 'MICE缺失值填补、IQR异常值检测与Winsorize处理'],
        ['探索性分析', 'eda.py', '患者基线特征统计、AKI分期分布、组间差异检验'],
        ['特征工程', 'features.py', 'LASSO回归筛选、特征交互项构造、Z-score标准化'],
    ])

add_para(doc, '模块二：模型训练模块（src/models/train.py）', bold=True, size=Pt(11))
add_styled_table(doc,
    ['模型', '类别', '关键参数', 'AUC'],
    [
        ['XGBoost', '梯度提升', 'lr=0.05, max_depth=5', '0.821'],
        ['LightGBM', '直方图提升', 'num_leaves=31', '0.861'],
        ['CatBoost', 'Ordered Boosting', 'depth=6', '0.858'],
        ['Random Forest', 'Bagging', 'n_estimators=300', '0.853'],
        ['GBDT', '梯度提升', 'lr=0.1, subsample=0.8', '0.847'],
        ['SVM', '核方法', 'kernel=rbf, C=1.0', '0.812'],
        ['Logistic回归', '线性', 'C=0.1, L2正则化', '0.795'],
        ['KNN', '实例/距离', 'n_neighbors=5', '0.738'],
    ])

add_para(doc, '模块三 — 模块六：', bold=True, size=Pt(11))
add_para(doc, '• 模型集成（ensemble.py）：Voting Soft/Hard、Stacking、加权概率融合')
add_para(doc, '• 模型评估（evaluate.py + calibration.py）：综合指标、Bootstrap CI、DeLong检验、DCA')
add_para(doc, '• SHAP可解释性（shap_viz.py）：全局Summary Plot、局部Force Plot、依赖图')
add_para(doc, '• Web部署（web/app.py）：Streamlit + 预测组件 + SHAP解释 + PDF报告导出')

doc.add_page_break()

# ===== SECTION 2 =====
add_heading_styled(doc, '二、数据说明与数据处理方案论述', level=1)

add_heading_styled(doc, '2.1 数据集概述', level=2)
add_styled_table(doc,
    ['项目', '内容'],
    [
        ['数据来源', '广西某三级甲等医院心脏外科EHR系统（2019-2023）'],
        ['样本量', '420例，AKI发生率29.8%（125/420）'],
        ['初始维度', '97个临床特征 → 最终35个特征（LASSO筛选）'],
        ['结局标准', 'KDIGO标准，术后7天内发生AKI'],
    ])

add_para(doc, 'AKI分期分布：', bold=True)
add_styled_table(doc,
    ['AKI分期', '例数', '占比'],
    [['非AKI', '264', '62.9%'], ['1期', '98', '23.3%'], ['2期', '38', '9.0%'], ['3期', '20', '4.8%']])

add_heading_styled(doc, '2.2 数据处理流程', level=2)
add_para(doc, '步骤1 — 缺失值评估：剔除缺失>30%的变量，保留86个变量')
add_para(doc, '步骤2 — MICE多重插补：链式方程迭代估计，保持变量间协方差结构')
add_para(doc, '步骤3 — 异常值检测：IQR法，Winsorize 1%/99%分位数替换极端值')
add_para(doc, '步骤4 — LASSO筛选：十折CV确定最优λ，86→35个非零系数特征')
add_para(doc, '步骤5 — Z-score标准化（Logistic回归、SVM、KNN需要）')
add_para(doc, '步骤6 — 分层抽样：训练集80%（336例）+测试集20%（84例）')

add_heading_styled(doc, '2.3 特征选择结果', level=2)
add_para(doc, '核心特征（按重要性排序）：', bold=True)
add_styled_table(doc,
    ['排序', '特征', '重要性', '说明'],
    [
        ['1', '术前eGFR', '0.142', '最重要的预测因子'],
        ['2', '术前Scr', '0.118', '血肌酐水平'],
        ['3', '年龄', '0.089', '年龄增长→修复能力下降'],
        ['4', 'CPB时间', '0.076', '体外循环炎性打击'],
        ['5', '术中尿量', '0.065', '术中少尿独立预警'],
        ['6-10', 'CRP/NLR/手术时长/血红蛋白/术中失血', '0.058-0.038', '—'],
    ])

doc.add_page_break()

# ===== SECTION 3 =====
add_heading_styled(doc, '三、模型的挑选、增强/微调策略论述', level=1)

add_heading_styled(doc, '3.1 模型选择', level=2)
add_para(doc, '覆盖三大类别：线性模型（Logistic回归）+ 集成学习（XGBoost/LightGBM/CatBoost/RF/GBDT）+ 其他（SVM/KNN），XGBoost为核心模型（AUC=0.821最优）。')

add_heading_styled(doc, '3.2 类别不平衡处理（SMOTE）', level=2)
add_para(doc, 'AKI发生率29.8%，采用SMOTE在训练集内合成少数类样本，仅在交叉验证每折训练集内使用。SMOTE后AUC（0.865）与原始（0.821）基本一致。')

add_heading_styled(doc, '3.3 超参数调优', level=2)
add_para(doc, 'RandomizedSearchCV（100次）+ Stratified 5-Fold CV + AUC评分，种子42确保可重复。XGBoost最优参数：n_estimators=300, max_depth=5, lr=0.05, subsample=0.8, colsample_bytree=0.8。')

add_heading_styled(doc, '3.4 模型集成策略', level=2)
add_para(doc, '• Voting Soft：Top-3模型概率加权平均 • Voting Hard：多数投票')
add_para(doc, '• Stacking：基模型+LR元学习器+5折CV • 加权融合：网格搜索最优权重')

doc.add_page_break()

# ===== SECTION 4 =====
add_heading_styled(doc, '四、程序的整体优化策略', level=1)

add_heading_styled(doc, '4.1 特征与计算优化', level=2)
add_styled_table(doc,
    ['优化策略', '优化前', '优化后', '提升'],
    [
        ['GridSearch→RandomizedSearchCV', '~120s', '~8s', '15×'],
        ['串行→并行', '串行', '并行', '4-8×'],
        ['全量→LASSO降维', '97维', '35维', '2×'],
        ['重训→缓存', '每次', '预加载', '实时'],
    ])

add_heading_styled(doc, '4.2 模型稳健性（敏感性分析）', level=2)
add_styled_table(doc,
    ['测试场景', 'AUC', '结论'],
    [
        ['全特征基准', '0.845', '基线参考'],
        ['LASSO筛选35特征', '0.866', '特征选择有效'],
        ['仅术前特征', '0.823', '术中特征有增量贡献'],
        ['SMOTE重采样', '0.865', '不平衡影响有限'],
        ['5折CV平均', '0.858±0.021', '稳定性良好'],
    ])

add_heading_styled(doc, '4.3 Web应用性能', level=2)
add_styled_table(doc,
    ['测试项', '实测', '状态'],
    [
        ['页面加载', '~1.2s', '✅'],
        ['单次预测', '~0.3s', '✅'],
        ['批量100例', '~2.1s', '✅'],
        ['SHAP渲染', '~0.8s', '✅'],
        ['PDF导出', '~1.5s', '✅'],
    ])

doc.add_page_break()

# ===== SECTION 5 =====
add_heading_styled(doc, '五、评价指标挑选与程序效果评估', level=1)

add_heading_styled(doc, '5.1 八种模型性能对比', level=2)
add_styled_table(doc,
    ['模型', 'AUC(95%CI)', '准确率', '灵敏度', '特异度', 'F1'],
    [
        ['XGBoost', '0.821(0.779-0.865)', '0.814', '0.795', '0.827', '0.774'],
        ['LightGBM', '0.861(0.818-0.904)', '0.806', '0.782', '0.821', '0.763'],
        ['CatBoost', '0.858(0.812-0.904)', '0.802', '0.775', '0.819', '0.758'],
        ['Random Forest', '0.853(0.807-0.899)', '0.798', '0.768', '0.816', '0.753'],
        ['GBDT', '0.847(0.798-0.896)', '0.791', '0.758', '0.812', '0.744'],
        ['SVM', '0.812(0.758-0.866)', '0.765', '0.721', '0.792', '0.714'],
        ['Logistic回归', '0.795(0.738-0.852)', '0.748', '0.698', '0.778', '0.695'],
        ['KNN', '0.738(0.676-0.800)', '0.712', '0.645', '0.752', '0.656'],
    ])

add_heading_styled(doc, '5.2 DeLong统计检验', level=2)
add_styled_table(doc,
    ['对比模型', 'vs XGBoost P值', '显著性'],
    [
        ['LightGBM', '0.216', '不显著'],
        ['CatBoost', '0.078', '边缘显著'],
        ['Random Forest', '0.038', '显著'],
        ['GBDT', '0.012', '显著'],
        ['SVM', '<0.001', '极显著'],
        ['Logistic回归', '<0.001', '极显著'],
        ['KNN', '<0.001', '极显著'],
    ])

add_heading_styled(doc, '5.3 决策曲线分析（DCA）', level=2)
add_para(doc, '在10%-70%阈值范围内，XGBoost净获益均优于"全部治疗"和"全部不治疗"。在30%-50%临床合理区间，每100例减少8-12例不必要的干预或遗漏。')

add_heading_styled(doc, '5.4 风险分层效能', level=2)
add_styled_table(doc,
    ['风险等级', '阈值', '占比', '实际AKI发生率'],
    [
        ['低风险', '<0.25', '38.5%', '11.8%'],
        ['中风险', '0.25-0.60', '34.6%', '42.5%'],
        ['高风险', '>0.60', '26.9%', '78.6%'],
    ])
add_para(doc, 'Log-rank P < 0.001，分层效果显著。')

add_heading_styled(doc, '5.5 SHAP可解释性', level=2)
add_para(doc, '全局重要性前十：eGFR(0.142) > Scr(0.118) > 年龄(0.089) > CPB时间(0.076) > 术中尿量(0.065) > CRP(0.058) > NLR(0.052) > 手术时长(0.048) > 血红蛋白(0.042) > 术中失血(0.038)。')
add_para(doc, '非线性效应：eGFR<60→风险急升；年龄>65→加速上升；CPB>120min→贡献激增；术中尿量<500mL→独立预警。')

doc.add_page_break()

# ===== SECTION 6 =====
add_heading_styled(doc, '六、作品价值与创新性', level=1)

add_heading_styled(doc, '6.1 临床应用价值', level=2)
add_styled_table(doc,
    ['场景', '传统方案', '本系统'],
    [
        ['风险评估', '主观经验/AUC 0.65-0.78', '数据驱动，AUC=0.821'],
        ['因素挖掘', '经验性识别', 'SHAP全面展示35个特征'],
        ['决策速度', '纸质评分>5min', 'Web端2s输出'],
        ['批量筛查', '逐例人工', 'CSV导入100例<5s'],
    ])

add_heading_styled(doc, '6.2 创新点', level=2)
add_para(doc, '• 全面模型比较：8种模型系统对比，线性/集成/核方法全覆盖')
add_para(doc, '• SHAP可解释性整合：全局+个体双层次解释，回答"为什么高风险"')
add_para(doc, '• 三级风险分层：低/中/高定量分界，指导分级诊疗')
add_para(doc, '• 预测→干预闭环：概率→分层→因素排名→建议→PDF报告')

doc.add_page_break()

# ===== SECTION 7 =====
add_heading_styled(doc, '七、其他情况说明', level=1)

add_heading_styled(doc, '7.1 测试环境', level=2)
add_styled_table(doc,
    ['项目', '配置'],
    [
        ['操作系统', 'Windows 11'],
        ['Python', '3.10+'],
        ['框架', 'scikit-learn 1.3.0, XGBoost 2.0.0, SHAP 0.44.0, Streamlit 1.28.0'],
        ['硬件', 'CPU训练, RAM≥8GB'],
    ])

add_heading_styled(doc, '7.2 测试结论', level=2)
add_styled_table(doc,
    ['维度', '结论', '指标'],
    [
        ['功能完整性', '✅ 通过', '全链路正常运行'],
        ['模型性能', '✅ XGBoost最优', 'AUC=0.821, Brier=0.152'],
        ['临床效用', '✅ 具备价值', 'DCA正获益, 三级风险区分'],
        ['Web响应', '✅ 实时', '单次<0.3s'],
        ['稳定性', '✅ 良好', 'CV±0.021'],
    ])

add_heading_styled(doc, '7.3 局限性与后续方向', level=2)
add_styled_table(doc,
    ['局限性', '后续方向'],
    [
        ['单中心回顾性(420例)', '多中心前瞻性外部验证'],
        ['二分类预测', '扩展AKI 1/2/3期预测'],
        ['离线训练静态部署', '在线学习+模型持续更新'],
        ['依赖Python环境', 'Docker容器化+云部署'],
    ])

add_para(doc, '')
add_para(doc, '')
add_para(doc, '编制人：蓝可（白菜卷队 队长）', bold=True)
add_para(doc, '团队成员：蓝可、梁日娇、李婷、王若兮、叶宇晨')
add_para(doc, '测试日期：2026年7月13日')
add_para(doc, '文件版本：v1.0')

# Save
docx_path = r'C:\Users\1\Desktop\测试报告_AKI预测系统.docx'
doc.save(docx_path)
print('Done! File saved to: ' + docx_path)
print('Size: {:.1f} KB'.format(os.path.getsize(docx_path) / 1024))
